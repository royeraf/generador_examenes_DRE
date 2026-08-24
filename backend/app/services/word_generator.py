"""
Servicio para generar documentos Word con el examen.
"""
from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from app.services.latex_word import add_latex_runs, set_cell_text, truncate_math_safe


def generar_examen_word(data: dict) -> BytesIO:
    """
    Genera un documento Word con el examen completo.
    
    Args:
        data: Diccionario con la estructura del examen generado.
        
    Returns:
        BytesIO: Buffer con el documento Word.
    """
    doc = Document()
    
    # Configurar márgenes
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    examen = data.get("examen", {})
    
    # Título del examen
    titulo = doc.add_heading(examen.get("titulo", "EXAMEN DE COMPRENSIÓN LECTORA"), 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Grado
    grado_para = doc.add_paragraph()
    grado_run = grado_para.add_run(f"Grado: {examen.get('grado', '')}")
    grado_run.bold = True
    grado_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Sección de datos del estudiante
    datos_table = doc.add_table(rows=1, cols=2)
    datos_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    cell1 = datos_table.rows[0].cells[0]
    cell1.text = "Apellidos y Nombres: _______________________________________"
    
    cell2 = datos_table.rows[0].cells[1]
    cell2.text = "Fecha: _______________________________________"
    
    doc.add_paragraph()
    
    # Instrucciones
    instrucciones_heading = doc.add_heading("INSTRUCCIONES", level=1)
    instrucciones_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    instrucciones_para = doc.add_paragraph()
    add_latex_runs(instrucciones_para, examen.get("instrucciones", "Lee atentamente el texto y responde las preguntas."))
    instrucciones_para.paragraph_format.space_after = Pt(12)

    # Lectura / situación problemática (MatSistem usa "situacion_problematica" en vez de "lectura")
    cuerpo_lectura = examen.get("lectura") or examen.get("situacion_problematica") or ""
    lectura_heading = doc.add_heading("LECTURA" if examen.get("lectura") else "SITUACIÓN PROBLEMÁTICA", level=1)

    lectura_para = doc.add_paragraph()
    add_latex_runs(lectura_para, cuerpo_lectura)
    lectura_para.paragraph_format.first_line_indent = Cm(1)
    lectura_para.paragraph_format.space_after = Pt(18)
    lectura_para.paragraph_format.line_spacing = 1.5
    
    # Preguntas
    preguntas_heading = doc.add_heading("PREGUNTAS", level=1)
    
    preguntas = examen.get("preguntas", [])
    for pregunta in preguntas:
        # Número y nivel de la pregunta
        numero = pregunta.get("numero", "")
        nivel = pregunta.get("nivel", "")
        enunciado = pregunta.get("enunciado", "")
        
        pregunta_para = doc.add_paragraph()
        pregunta_run = pregunta_para.add_run(f"{numero}. ")
        pregunta_run.bold = True

        if nivel:
            pregunta_para.add_run(f"[{nivel}] ")
        add_latex_runs(pregunta_para, enunciado)
        pregunta_para.paragraph_format.space_before = Pt(12)

        # Opciones
        opciones = pregunta.get("opciones", [])
        for opcion in opciones:
            opcion_para = doc.add_paragraph()
            opcion_para.paragraph_format.left_indent = Cm(1)
            letra = opcion.get("letra", "")
            texto = opcion.get("texto", "")
            opcion_para.add_run(f"    {letra}) ")
            add_latex_runs(opcion_para, texto)

        doc.add_paragraph()
    
    # Separador antes de la tabla de respuestas
    doc.add_page_break()
    
    # Tabla de respuestas (para el docente)
    tabla_heading = doc.add_heading("TABLA DE RESPUESTAS (PARA EL DOCENTE)", level=1)
    tabla_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tabla_respuestas = examen.get("tabla_respuestas", [])
    
    if tabla_respuestas:
        # Crear tabla
        table = doc.add_table(rows=len(tabla_respuestas) + 1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Encabezados
        header_cells = table.rows[0].cells
        headers = ["# Pregunta", "Desempeño", "Nivel", "Respuesta", "Justificación"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            header_cells[i].paragraphs[0].runs[0].bold = True
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Datos
        for i, fila in enumerate(tabla_respuestas):
            row_cells = table.rows[i + 1].cells
            set_cell_text(row_cells[0], str(fila.get("pregunta", "")), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row_cells[1], truncate_math_safe(fila.get("desempeno", ""), 50))
            set_cell_text(row_cells[2], fila.get("nivel", ""), align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(
                row_cells[3],
                fila.get("respuesta_correcta") or fila.get("respuesta_esperada", ""),
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            set_cell_text(row_cells[4], fila.get("justificacion", ""))
    
    # Guardar en buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer
