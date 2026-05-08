"""
Servicio para extraer texto de archivos PDF y Word.
Incluye validaciones de seguridad multicapa para prevenir archivos maliciosos.
"""
import fitz  # PyMuPDF
from docx import Document
from fastapi import UploadFile, HTTPException
import io
import re
import unicodedata
from typing import Tuple
import zipfile
import logging


logger = logging.getLogger(__name__)


class FileExtractionService:

    MAX_FILE_SIZE_BYTES = 10 * 1024 * 1024
    ALLOWED_EXTENSIONS = {"pdf", "docx", "doc"}

    MAGIC_BYTES: dict[str, list[bytes]] = {
        "pdf":  [b"%PDF"],
        "docx": [b"PK\x03\x04"],
        "doc":  [b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"],
    }

    PDF_DANGEROUS_PATTERNS = [
        rb"/JavaScript",
        rb"/JS\b",
        rb"/AA\b",
        rb"/OpenAction",
        rb"/Launch",
        rb"/EmbeddedFile",
        rb"/RichMedia",
        rb"/XFA",
        rb"eval\s*\(",
        rb"<script",
    ]

    DOCX_DANGEROUS_PATTERNS = [
        r"<\s*script",
        r"javascript\s*:",
        r"vbscript\s*:",
        r"on\w+\s*=",
        r"<\s*object\b",
        r"<\s*embed\b",
        r"<\s*iframe\b",
        r"macroEnabled",
        r"w:macros",
    ]

    DOCX_XML_MEMBERS = [
        "word/document.xml",
        "word/settings.xml",
        "word/webSettings.xml",
        "[Content_Types].xml",
        "_rels/.rels",
    ]

    # ── Seguridad ──────────────────────────────────────────────────────────────

    def _sanitize_filename(self, filename: str) -> str:
        filename = unicodedata.normalize("NFKD", filename)
        filename = filename.replace("..", "").replace("/", "").replace("\\", "")
        filename = re.sub(r"[^\w\s.\-]", "_", filename)
        filename = filename[:200]
        return filename.strip() or "archivo_sin_nombre"

    def _validate_magic_bytes(self, content: bytes, extension: str) -> None:
        expected = self.MAGIC_BYTES.get(extension, [])
        if not expected:
            raise HTTPException(400, f"Extensión '{extension}' no tiene firma conocida.")
        for sig in expected:
            if content.startswith(sig):
                return
        raise HTTPException(
            status_code=400,
            detail=(
                f"El archivo no es un {extension.upper()} válido. "
                "El contenido real no coincide con la extensión declarada. "
                "Posible archivo malicioso o renombrado."
            ),
        )

    def _scan_pdf_for_threats(self, content: bytes) -> None:
        for pattern in self.PDF_DANGEROUS_PATTERNS:
            if re.search(pattern, content, re.IGNORECASE):
                logger.warning("PDF rechazado: patrón peligroso detectado: %s", pattern)
                raise HTTPException(
                    status_code=400,
                    detail=(
                        "El archivo PDF contiene elementos potencialmente peligrosos "
                        "(JavaScript, acciones automáticas u objetos embebidos). "
                        "Por seguridad, el archivo fue rechazado."
                    ),
                )

    def _scan_docx_for_threats(self, content: bytes) -> None:
        try:
            with zipfile.ZipFile(io.BytesIO(content)) as zf:
                member_names = zf.namelist()

                dangerous_embedded = [
                    name for name in member_names
                    if name.lower().endswith(
                        (".exe", ".bat", ".cmd", ".ps1", ".vbs", ".js", ".jar", ".sh")
                    )
                ]
                if dangerous_embedded:
                    logger.warning("DOCX rechazado: archivos ejecutables embebidos: %s", dangerous_embedded)
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            "El archivo Word contiene archivos ejecutables embebidos. "
                            "Por seguridad, el archivo fue rechazado."
                        ),
                    )

                for member in self.DOCX_XML_MEMBERS:
                    if member in member_names:
                        try:
                            xml_content = zf.read(member).decode("utf-8", errors="replace")
                            for pattern in self.DOCX_DANGEROUS_PATTERNS:
                                if re.search(pattern, xml_content, re.IGNORECASE):
                                    logger.warning(
                                        "DOCX rechazado: patrón peligroso '%s' en %s", pattern, member
                                    )
                                    raise HTTPException(
                                        status_code=400,
                                        detail=(
                                            "El archivo Word contiene scripts o macros potencialmente peligrosos. "
                                            "Por seguridad, el archivo fue rechazado."
                                        ),
                                    )
                        except HTTPException:
                            raise
                        except Exception:
                            pass

                macro_files = [
                    name for name in member_names
                    if name.lower().startswith("word/vba") or name.lower().endswith(".bin")
                ]
                if macro_files:
                    logger.warning("DOCX rechazado: macros VBA detectadas: %s", macro_files)
                    raise HTTPException(
                        status_code=400,
                        detail=(
                            "El archivo Word contiene macros VBA. "
                            "Por seguridad, solo se aceptan documentos sin macros (.docx estándar)."
                        ),
                    )

        except HTTPException:
            raise
        except zipfile.BadZipFile:
            raise HTTPException(400, "El archivo .docx está corrupto o no es un documento Word válido.")
        except Exception as e:
            logger.error("Error al escanear DOCX: %s", e)
            raise HTTPException(400, "No se pudo verificar la seguridad del archivo Word.")

    # ── Extracción de texto ────────────────────────────────────────────────────

    def _extract_text_from_pdf(self, content: bytes) -> str:
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = [page.get_text() for page in doc if page.get_text().strip()]
            doc.close()
            return "\n".join(text_parts)
        except Exception as e:
            raise HTTPException(400, f"Error al procesar PDF: {e}")

    def _extract_text_from_docx(self, content: bytes) -> str:
        try:
            doc = Document(io.BytesIO(content))
            text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)
            return "\n".join(text_parts)
        except Exception as e:
            raise HTTPException(400, f"Error al procesar Word: {e}")

    # ── Pipeline público ───────────────────────────────────────────────────────

    async def extract_text_from_file(self, file: UploadFile) -> Tuple[str, dict]:
        """
        Pipeline de seguridad multicapa:
        1. Validar extensión  2. Sanitizar nombre  3. Verificar tamaño
        4. Verificar magic bytes  5. Escanear amenazas  6. Extraer texto
        """
        filename = file.filename or "sin_nombre"
        extension = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if extension not in self.ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=(
                    f"Tipo de archivo no permitido: '.{extension}'. "
                    f"Solo se aceptan: {', '.join(sorted(self.ALLOWED_EXTENSIONS))}"
                ),
            )

        safe_filename = self._sanitize_filename(filename)
        content = await file.read()

        if len(content) == 0:
            raise HTTPException(400, "El archivo está vacío.")
        if len(content) > self.MAX_FILE_SIZE_BYTES:
            size_mb = len(content) / (1024 * 1024)
            raise HTTPException(
                status_code=413,
                detail=(
                    f"El archivo '{safe_filename}' pesa {size_mb:.1f} MB. "
                    f"El tamaño máximo permitido es {self.MAX_FILE_SIZE_BYTES // (1024 * 1024)} MB."
                ),
            )

        self._validate_magic_bytes(content, extension)

        if extension == "pdf":
            self._scan_pdf_for_threats(content)
        elif extension == "docx":
            self._scan_docx_for_threats(content)

        if extension == "pdf":
            text = self._extract_text_from_pdf(content)
        else:
            text = self._extract_text_from_docx(content)

        text = text.strip()
        if not text:
            raise HTTPException(
                status_code=400,
                detail=(
                    "No se pudo extraer texto del archivo. "
                    "Verifique que el archivo contenga texto legible (no sea una imagen escaneada)."
                ),
            )

        metadata = {
            "filename": safe_filename,
            "filename_original": filename,
            "extension": extension,
            "size_bytes": len(content),
            "size_kb": round(len(content) / 1024, 1),
            "caracteres": len(text),
            "palabras": len(text.split()),
            "lineas": text.count("\n") + 1,
        }

        logger.info(
            "Archivo procesado correctamente: %s (%s KB, %s palabras)",
            safe_filename, metadata["size_kb"], metadata["palabras"],
        )

        return text, metadata


file_extraction_service = FileExtractionService()
